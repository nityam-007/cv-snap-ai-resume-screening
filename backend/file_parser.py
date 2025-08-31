"""
CV Snap - File Parsing Utilities
Handles PDF and DOCX file extraction
"""

import PyPDF2
from docx import Document
import io
import re
from typing import Optional
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentParser:
    """Utility class for parsing PDF and DOCX files"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove extra whitespace and normalize line breaks
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        text = text.strip()
        return text
    
    @staticmethod
    def parse_pdf(file_content: bytes) -> Optional[str]:
        """
        Extract text from PDF file
        
        Args:
            file_content (bytes): PDF file content as bytes
            
        Returns:
            Optional[str]: Extracted text or None if parsing fails
        """
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            extracted_text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    extracted_text += page_text + "\n"
                    logger.info(f"Successfully extracted text from PDF page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                    continue
            
            if not extracted_text.strip():
                logger.error("No text could be extracted from PDF")
                return None
                
            cleaned_text = DocumentParser.clean_text(extracted_text)
            logger.info(f"Successfully parsed PDF. Extracted {len(cleaned_text)} characters")
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            return None
    
    @staticmethod
    def parse_docx(file_content: bytes) -> Optional[str]:
        """
        Extract text from DOCX file
        
        Args:
            file_content (bytes): DOCX file content as bytes
            
        Returns:
            Optional[str]: Extracted text or None if parsing fails
        """
        try:
            doc = Document(io.BytesIO(file_content))
            extracted_text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            extracted_text += cell.text + " "
                    extracted_text += "\n"
            
            if not extracted_text.strip():
                logger.error("No text could be extracted from DOCX")
                return None
                
            cleaned_text = DocumentParser.clean_text(extracted_text)
            logger.info(f"Successfully parsed DOCX. Extracted {len(cleaned_text)} characters")
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return None
    
    @staticmethod
    def parse_document(filename: str, file_content: bytes) -> Optional[str]:
        """
        Parse document based on file extension
        
        Args:
            filename (str): Name of the file
            file_content (bytes): File content as bytes
            
        Returns:
            Optional[str]: Extracted text or None if parsing fails
        """
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return DocumentParser.parse_pdf(file_content)
        elif file_extension == 'docx':
            return DocumentParser.parse_docx(file_content)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return None
    
    @staticmethod
    def extract_basic_info(text: str) -> dict:
        """
        Extract basic information from resume text using regex patterns
        
        Args:
            text (str): Resume text
            
        Returns:
            dict: Basic extracted information
        """
        info = {
            'emails': [],
            'phones': [],
            'linkedin': None,
            'github': None
        }
        
        try:
            # Email pattern
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            info['emails'] = re.findall(email_pattern, text)
            
            # Phone pattern (various formats)
            phone_pattern = r'(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
            info['phones'] = re.findall(phone_pattern, text)
            
            # LinkedIn pattern
            linkedin_pattern = r'linkedin\.com/in/[\w-]+'
            linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
            if linkedin_match:
                info['linkedin'] = linkedin_match.group()
            
            # GitHub pattern
            github_pattern = r'github\.com/[\w-]+'
            github_match = re.search(github_pattern, text, re.IGNORECASE)
            if github_match:
                info['github'] = github_match.group()
                
        except Exception as e:
            logger.warning(f"Error extracting basic info: {str(e)}")
        
        return info

# Example usage and testing
if __name__ == "__main__":
    # Test the parser with sample content
    parser = DocumentParser()
    
    # Test text cleaning
    sample_text = "   This   is    a  test   \n\n\n   with   lots    of   spaces   "
    cleaned = parser.clean_text(sample_text)
    print(f"Cleaned text: '{cleaned}'")
    
    # Test basic info extraction
    sample_resume = """
    John Doe
    john.doe@email.com
    Phone: +1 (555) 123-4567
    LinkedIn: linkedin.com/in/johndoe
    GitHub: github.com/johndoe
    
    Software Engineer with 5 years of experience...
    """
    
    basic_info = parser.extract_basic_info(sample_resume)
    print("Extracted basic info:", basic_info)
